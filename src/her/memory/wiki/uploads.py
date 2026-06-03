"""Turn an uploaded file into something Claude Opus can ingest into the wiki.

Each supported file type maps to one of two shapes:

* **text** — for ``.txt``/``.md`` (read as-is) and ``.docx`` (paragraphs pulled
  out with ``python-docx``). The text becomes the wiki source body.
* **attachments** — for ``.pdf`` and images (``.jpg``/``.png``), which are sent
  to Opus natively as Anthropic content blocks (a ``document`` block for PDFs,
  an ``image`` block for pictures) so the model reads/sees them directly.

``extract_source`` returns ``{"label", "text", "attachments"}`` with exactly one
of ``text``/``attachments`` populated. The server validates the extension with
``is_allowed`` before saving, so ``extract_source`` only sees allowed types.
"""
from __future__ import annotations

import base64
import logging
import secrets
from dataclasses import dataclass
from pathlib import Path

log = logging.getLogger(__name__)

# Allowed upload extensions → the media type used when handing the file to Opus.
TEXT_EXTS = {".txt", ".md"}
DOCX_EXTS = {".docx"}
PDF_EXTS = {".pdf"}
IMAGE_MEDIA = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png"}

ALLOWED_EXTS = TEXT_EXTS | DOCX_EXTS | PDF_EXTS | set(IMAGE_MEDIA)


def is_allowed(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXTS


def _b64(path: Path) -> str:
    return base64.standard_b64encode(path.read_bytes()).decode("ascii")


def _docx_text(path: Path) -> str:
    """Extract paragraph + table text from a .docx via python-docx."""
    try:
        import docx  # noqa: PLC0415  (optional dependency, imported lazily)
    except ImportError as e:
        raise RuntimeError(
            "the 'python-docx' package is not installed — run "
            "`pip install python-docx` to ingest .docx files"
        ) from e
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_source(path: str | Path) -> dict:
    """Return ``{"label", "text", "attachments"}`` for an uploaded file.

    Exactly one of ``text`` / ``attachments`` is set. Raises ``ValueError`` for
    an unsupported extension.
    """
    path = Path(path)
    ext = path.suffix.lower()
    label = path.name
    if ext in TEXT_EXTS:
        return {"label": label, "text": path.read_text(encoding="utf-8", errors="replace"),
                "attachments": None}
    if ext in DOCX_EXTS:
        return {"label": label, "text": _docx_text(path), "attachments": None}
    if ext in PDF_EXTS:
        block = {
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": _b64(path),
            },
        }
        return {"label": label, "text": None, "attachments": [block]}
    if ext in IMAGE_MEDIA:
        block = {
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": IMAGE_MEDIA[ext],
                "data": _b64(path),
            },
        }
        return {"label": label, "text": None, "attachments": [block]}
    raise ValueError(f"unsupported file type: {ext or '(none)'}")


# ── Pending uploads ────────────────────────────────────────────────────────
# A file is saved and registered here on upload, then resolved by a follow-up
# decision: keep (ingest into the wiki) or discard (read once, then delete).
# Process-wide and in-memory — a pending upload that is never resolved simply
# lingers on disk under uploads_path until the next manual cleanup.


@dataclass
class PendingUpload:
    id: str
    path: Path
    label: str


class PendingUploadStore:
    def __init__(self) -> None:
        self._items: dict[str, PendingUpload] = {}

    def add(self, path: str | Path, label: str) -> PendingUpload:
        item = PendingUpload(id=secrets.token_hex(8), path=Path(path), label=label)
        self._items[item.id] = item
        return item

    def get(self, upload_id: str) -> PendingUpload | None:
        return self._items.get(upload_id)

    def pop(self, upload_id: str) -> PendingUpload | None:
        return self._items.pop(upload_id, None)


# Process-wide singleton.
pending = PendingUploadStore()
