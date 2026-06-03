"""Filesystem-layer tests for Cowork skills and the knowledge wiki.

These cover only the parts that don't touch the Anthropic API: the on-disk
Agent-Skill store and the wiki store. The LLM engine is exercised indirectly
via its graceful-degradation path (no credential -> friendly string).
"""
from __future__ import annotations

import asyncio

from her.core.usage import UsageTracker
from her.cowork.skills_store import CoworkSkillStore, slugify
from her.memory.wiki.engine import WikiEngine
from her.memory.wiki.store import WikiStore
from her.memory.wiki.uploads import PendingUploadStore, extract_source, is_allowed


# ---------- CoworkSkillStore ---------------------------------------------


def test_skill_write_list_read_delete(tmp_path):
    store = CoworkSkillStore(tmp_path / "skills")
    slug = store.write_skill(
        "Send Weekly Report",
        "Use when the user asks to compile the weekly report",
        "# Weekly Report\n\nDo the thing.",
    )
    assert slug == "send-weekly-report"

    skills = store.list_skills()
    assert skills == [{
        "slug": "send-weekly-report",
        "name": "send-weekly-report",
        "description": "Use when the user asks to compile the weekly report",
    }]

    md = store.read_skill(slug)
    assert md.startswith("---")
    assert "name: send-weekly-report" in md
    assert "# Weekly Report" in md

    assert store.delete_skill(slug) is True
    assert store.list_skills() == []
    assert store.delete_skill(slug) is False


def test_skill_description_is_single_line(tmp_path):
    store = CoworkSkillStore(tmp_path / "skills")
    store.write_skill("X", "line one\nline two", "# X\n\nbody")
    md = store.read_skill("x")
    # The one-line description must not break the frontmatter block.
    assert "description: line one line two" in md


def test_slugify():
    assert slugify("Hello World!") == "hello-world"
    assert slugify("  ") == "skill"


# ---------- WikiStore -----------------------------------------------------


def test_wiki_init_creates_scaffold(tmp_path):
    store = WikiStore(tmp_path / "wiki")
    store.ensure_init()
    assert (tmp_path / "wiki" / "CLAUDE.md").exists()
    assert (tmp_path / "wiki" / "index.md").exists()
    assert (tmp_path / "wiki" / "log.md").exists()
    assert (tmp_path / "wiki" / "pages").is_dir()


def test_wiki_page_roundtrip_and_index(tmp_path):
    store = WikiStore(tmp_path / "wiki")
    store.ensure_init()
    slug = store.write_page("costco", "Costco", "A wholesale retailer", "# Costco\n\nSee [[membership]].")
    assert slug == "costco"

    pages = store.list_pages()
    assert pages[0]["slug"] == "costco"
    assert pages[0]["title"] == "Costco"
    assert pages[0]["summary"] == "A wholesale retailer"

    body = store.page_body("costco")
    assert body.startswith("# Costco")
    assert "[[membership]]" in body

    store.rebuild_index()
    index = store.read_index()
    assert "[Costco](pages/costco.md)" in index
    assert "A wholesale retailer" in index

    assert store.delete_page("costco") is True
    assert store.list_pages() == []


def test_wiki_append_log(tmp_path):
    store = WikiStore(tmp_path / "wiki")
    store.ensure_init()
    store.append_log("ingest", "Costco 10-K", "added revenue figures")
    text = (tmp_path / "wiki" / "log.md").read_text(encoding="utf-8")
    assert "ingest | Costco 10-K" in text
    assert "added revenue figures" in text


# ---------- WikiEngine graceful degradation -------------------------------


def test_wiki_engine_no_credential(monkeypatch):
    # With no Anthropic credential, every LLM op returns a friendly string and
    # never raises (so a session keeps running). This path short-circuits in
    # ``_ready()`` before touching the filesystem, so no store setup is needed.
    from her.cowork.client import client as cowork

    monkeypatch.setattr(cowork, "is_configured", lambda: False)
    eng = WikiEngine()

    out = asyncio.run(eng.query("anything"))
    assert "credential" in out.lower()
    out = asyncio.run(eng.ingest("some text", "src"))
    assert "credential" in out.lower()
    out = asyncio.run(eng.ingest_file("/tmp/whatever.txt", "src"))
    assert "credential" in out.lower()
    out = asyncio.run(eng.read_file("/tmp/whatever.txt", "src"))
    assert "credential" in out.lower()


# ---------- Upload extractor ----------------------------------------------


def test_extract_source_text_and_markdown(tmp_path):
    p = tmp_path / "note.txt"
    p.write_text("hello world", encoding="utf-8")
    r = extract_source(p)
    assert r["text"] == "hello world" and r["attachments"] is None
    assert r["label"] == "note.txt"


def test_extract_source_docx(tmp_path):
    import docx

    doc = docx.Document()
    doc.add_paragraph("Para one")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    p = tmp_path / "report.docx"
    doc.save(str(p))
    r = extract_source(p)
    assert "Para one" in r["text"] and "A | B" in r["text"]
    assert r["attachments"] is None


def test_extract_source_pdf_and_image_become_attachments(tmp_path):
    pdf = tmp_path / "f.pdf"
    pdf.write_bytes(b"%PDF-1.4 minimal")
    r = extract_source(pdf)
    assert r["text"] is None
    assert r["attachments"][0]["type"] == "document"
    assert r["attachments"][0]["source"]["media_type"] == "application/pdf"

    from PIL import Image

    img = tmp_path / "pic.png"
    Image.new("RGB", (4, 4), (255, 0, 0)).save(str(img))
    r = extract_source(img)
    assert r["attachments"][0]["type"] == "image"
    assert r["attachments"][0]["source"]["media_type"] == "image/png"


def test_extract_source_rejects_unsupported(tmp_path):
    assert is_allowed("a.pdf") and not is_allowed("a.zip")
    p = tmp_path / "a.zip"
    p.write_bytes(b"PK")
    try:
        extract_source(p)
        raise AssertionError("expected ValueError")
    except ValueError as e:
        assert ".zip" in str(e)


def test_pending_upload_store_add_get_pop(tmp_path):
    store = PendingUploadStore()
    item = store.add(tmp_path / "x.txt", "x.txt")
    assert store.get(item.id) is item
    assert store.pop(item.id) is item
    assert store.get(item.id) is None
    assert store.pop(item.id) is None


# ---------- Anthropic cost accounting -------------------------------------


def test_record_anthropic_cost_and_total():
    u = UsageTracker()
    u.reset(model="gpt-realtime-mini")
    # Pretend the OpenAI side spent nothing; only Claude is charged.
    u.record_anthropic(
        {"input_tokens": 1000, "output_tokens": 500,
         "cache_read_input_tokens": 0, "cache_creation_input_tokens": 0},
        "claude-opus-4-8",
    )
    snap = u.snapshot()
    # opus-4-8: 1000*5/1e6 + 500*25/1e6 = 0.0175
    assert snap["anthropic"]["cost_usd"] == 0.0175
    assert snap["anthropic"]["requests"] == 1
    assert snap["cost_usd"] == 0.0  # OpenAI bucket untouched
    assert snap["cost_total_usd"] == 0.0175  # combined


def test_record_anthropic_object_with_cache():
    u = UsageTracker()
    u.reset(model="gpt-realtime-mini")

    class _Usage:
        input_tokens = 0
        output_tokens = 0
        cache_read_input_tokens = 10000  # billed at 0.1x input
        cache_creation_input_tokens = 0

    u.record_anthropic(_Usage(), "claude-sonnet-4-6")
    # sonnet input $3/1M -> cache read 0.3/1M -> 10000 * 0.3/1e6 = 0.000003 -> 0.0 at 5dp? no: 0.000003
    assert u.snapshot()["anthropic"]["cost_usd"] == round(10000 * (3.0 * 0.1) / 1_000_000, 5)
    # None is a no-op (never raises).
    u.record_anthropic(None)
    assert u.snapshot()["anthropic"]["requests"] == 1
